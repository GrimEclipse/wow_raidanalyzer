import json
from collections import Counter
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt


DATA = Path("tmp/crown_fight30_audit.json")
OUT = Path("output/宇宙之冕_Fight30_量化明细.docx")


def value(item):
    if item is None:
        return "-"
    if isinstance(item, bool):
        return "是" if item else "否"
    return str(item)


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header
    for row in rows:
        cells = table.add_row().cells
        for index, item in enumerate(row):
            cells[index].text = value(item)
    if widths:
        for row in table.rows:
            for index, width in enumerate(widths):
                row.cells[index].width = Cm(width)
    doc.add_paragraph()
    return table


def main():
    data = json.loads(DATA.read_text(encoding="utf-8"))
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.left_margin = section.right_margin = Cm(1.2)
    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal.font.size = Pt(9)

    title = doc.add_heading("宇宙之冕单场量化明细", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(
        f"报告 {data['meta']['reportID']} · Fight {data['meta']['fightID']} · "
        f"{'击杀' if data['fight']['kill'] else '未击杀'} · 时长 {data['fight']['duration']}"
    )

    doc.add_heading("1. 场次概览", level=1)
    phases = data.get("phase") or {}
    add_table(doc, ["项目", "数值"], [
        ("拉弓组数", len(data.get("bowGroups") or [])),
        ("放水轮数 / 落点数", f"{len(data.get('waterEvents') or [])} / {len(data.get('waterDrops') or [])}"),
        ("P1 银锋箭轮数", sum(1 for row in data.get("silverArrows") or [] if row.get("phase") == "P1")),
        ("P2 银锋弹射轮数", sum(1 for row in data.get("silverArrows") or [] if row.get("phase") == "P2")),
        ("P1.5 开始", phases.get("p15Start")), ("P2 开始", phases.get("p2Start")),
        ("P2.5 开始", phases.get("p25Start")), ("P3 开始", phases.get("p3Start")),
    ])

    doc.add_heading("2. 死亡明细与前五死", level=1)
    deaths = data.get("deathDetails") or []
    add_table(doc, ["死亡顺位", "时间", "阶段", "玩家", "致死技能ID"], [
        (index, row.get("time"), row.get("phase"), row.get("player"), row.get("abilityID"))
        for index, row in enumerate(deaths, start=1)
    ] or [("-", "-", "-", "本场无死亡", "-")])
    doc.add_paragraph("放弃判定：当前单场审计没有可靠的团队主动放弃事件；本表只统计日志实际死亡，不把战斗结束自动解释为放弃。")

    doc.add_heading("3. 空虚之握 / 崩裂空无逐轮明细", level=1)
    bow_rows = []
    for group in data.get("bowGroups") or []:
        for player in group.get("players") or []:
            healing = player.get("allHealing") or {}
            breakdown = "；".join(f"{item['healer']} {item['amount']}" for item in healing.get("healingByHealer") or [])
            bow_rows.append((
                group.get("index"), group.get("phase"), group.get("fireTime"), player.get("player"),
                player.get("applyTime"), player.get("fadeTime"), healing.get("totalHealing", 0), breakdown or "无",
                player.get("lastSecondMovementYards"), player.get("isSnapAiming"),
                player.get("actualPhantomHitCount", 0),
                ",".join(map(str, player.get("resolvedPhantomInstances") or [])) or "-",
                player.get("resolutionEvidence"),
            ))
    add_table(doc,
        ["轮次", "阶段", "结算", "玩家", "apply", "remove", "全程治疗", "治疗者拆分", "末1秒位移(码)", "甩狙", "消失幻影数", "实例", "证据"],
        bow_rows)

    doc.add_heading("4. 空虚之握死亡前6秒治疗", level=1)
    void_deaths = data.get("voidDeaths") or []
    add_table(doc, ["时间", "玩家", "幻影数", "治疗总量", "治疗者拆分", "幻影≥4豁免"], [
        (row.get("time"), row.get("player"), row.get("activePhantomCount"), row.get("totalHealing"),
         "；".join(f"{item['healer']} {item['amount']}" for item in row.get("healingByHealer") or []), row.get("exemptByPhantoms"))
        for row in void_deaths
    ] or [("-", "本场无此类死亡", "-", 0, "-", "-")])

    doc.add_heading("5. 放水逐落点明细", level=1)
    add_table(doc, ["阶段", "轮次", "时间", "玩家", "X(码)", "Y(码)", "离组(码)", "偏离>8码", "成圈时间"], [
        (drop.get("phase"), next((event.get("index") for event in data.get("waterEvents") or [] if any(item.get("id") == drop.get("id") for item in event.get("drops") or [])), "-"),
         drop.get("time"), drop.get("player"), (drop.get("position") or {}).get("yardX"), (drop.get("position") or {}).get("yardY"),
         drop.get("distanceFromGroupYards"), drop.get("isOutlier", False),
         next((event.get("time") for event in data.get("waterEvents") or [] if event.get("timeMs") == drop.get("maturesAtMs")), "下一轮/未成圈"))
        for drop in data.get("waterDrops") or []
    ])

    doc.add_heading("6. P1 银锋箭", level=1)
    add_table(doc, ["轮次", "射出时间", "点名玩家", "点名坐标", "误伤目标", "误伤次数"], [
        (row.get("index"), row.get("time"), " / ".join(row.get("markedPlayers") or []),
         "；".join(f"{item['player']}({(item.get('position') or {}).get('yardX')},{(item.get('position') or {}).get('yardY')})" for item in row.get("markedPlayerPositions") or []),
         " / ".join(hit.get("target", "?") for hit in row.get("actualHits") or []), len(row.get("actualHits") or []))
        for row in data.get("silverArrows") or [] if row.get("phase") == "P1"
    ])

    doc.add_heading("7. P2 射影与银锋弹射", level=1)
    add_table(doc, ["轮次", "结算时间", "场上幻影实例", "点名玩家", "几何命中", "实际归因"], [
        (group.get("index"), group.get("fireTime"), ",".join(str(item.get("sourceInstance")) for item in group.get("phantoms") or []),
         " / ".join(player.get("player") for player in group.get("players") or []),
         "；".join(f"{player.get('player')}→{','.join(str(hit.get('phantom')) for hit in player.get('predictedPhantomHits') or []) or '空'}" for player in group.get("players") or []),
         "；".join(f"{player.get('player')}→{','.join(map(str, player.get('resolvedPhantomInstances') or [])) or '无'}" for player in group.get("players") or []))
        for group in data.get("bowGroups") or [] if group.get("phase") == "P2"
    ])
    add_table(doc, ["轮次", "时间", "来源幻影", "点名玩家", "是否消能", "失败玩家"], [
        (row.get("index"), row.get("time"), ",".join(str(item.get("sourceInstance")) for item in row.get("sourcePhantoms") or []),
         " / ".join(row.get("markedPlayers") or []), row.get("bossEnergyDrained"), " / ".join(row.get("failedPlayers") or []))
        for row in data.get("silverArrows") or [] if row.get("phase") == "P2"
    ])

    doc.add_heading("8. P3 传送门、能量与场地污染", level=1)
    add_table(doc, ["事件", "时间", "Boss能量", "能量上限", "全团坐标数", "Boss坐标数"], [
        ("宇宙传送门" if row.get("eventType") == "portal" else "噬灭宇宙", row.get("time"),
         (row.get("bossEnergy") or {}).get("amount"), (row.get("bossEnergy") or {}).get("max"),
         len((row.get("snapshot") or {}).get("players") or []), len((row.get("snapshot") or {}).get("bosses") or []))
        for row in data.get("p3Events") or []
    ])
    add_table(doc, ["噬灭时间(ms)", "彻底污染时间(ms)", "板块", "Boss坐标"], [
        (row.get("castTimeMs"), row.get("activeTimeMs"), row.get("platform"),
         f"{(row.get('bossPosition') or {}).get('yardX')},{(row.get('bossPosition') or {}).get('yardY')}")
        for row in data.get("p3Contaminations") or []
    ])

    doc.add_heading("9. 当前尚不能可靠量化的字段", level=1)
    for text in (
        "主动放弃：日志没有稳定的放弃事件，不能仅凭战斗结束推断。",
        "P1 银锋箭是否开个人减伤：需要另建减伤技能白名单和伤害窗口。",
        "P1.5 吃球未死：当前单场审计尚未抓取对应球体伤害事件。",
        "P3 拉线第1/2/3棒及拉断间隔：需要确认连线 debuff ID 与 remove 事件后再量化。",
        "崩裂空无害死人责任：目前可输出伤害与实例/几何证据，但无法对证据冲突样本强制归责。",
    ):
        doc.add_paragraph(text, style="List Bullet")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT.resolve())


if __name__ == "__main__":
    main()
